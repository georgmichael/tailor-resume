#!/usr/bin/env python3
"""
Apply rewritten bullet text back into a .docx, preserving formatting.

Reads a JSON file of {"<paragraph_id>": "new text", ...} (ids match those from
extract_bullets.py) and writes the new text into the corresponding paragraphs
WITHOUT touching paragraph style, numbering, indentation, or fonts.

Formatting-preservation strategy:
  - We never add or remove paragraphs (ids stay valid).
  - For each target paragraph we keep run #0 and write the new text into it,
    inheriting that run's character formatting (font, size, bold, color).
  - Remaining runs are emptied (not deleted) so the XML structure is untouched.
  - If the bullet uses a *manual* leading marker char (•, -, etc.) rather than
    Word's auto-numbering, we preserve that exact prefix + spacing.

Usage:
    python apply_rewrites.py <resume.docx> <rewrites.json> --out <tailored.docx>
"""
import argparse
import json
import re
import sys

from docx import Document

MARKER_RE = re.compile(r"^(\s*[•\-–—▪◦‣*]\s+)")


def split_marker(original_text: str):
    """Return (prefix, body) where prefix is a manual bullet marker if present."""
    m = MARKER_RE.match(original_text)
    if m:
        return m.group(1), original_text[m.end():]
    return "", original_text


def set_paragraph_text(paragraph, new_text: str):
    """Replace paragraph text, preserving run-0 formatting and the bullet marker."""
    original = paragraph.text
    prefix, _ = split_marker(original)

    # Strip any marker the rewrite may have accidentally included so we don't
    # double up; we re-apply the original prefix.
    _, new_body = split_marker(new_text)
    final_text = f"{prefix}{new_body}" if prefix else new_body

    runs = paragraph.runs
    if not runs:
        # No runs (rare) -> just add one; inherits paragraph style.
        paragraph.add_run(final_text)
        return

    runs[0].text = final_text
    for r in runs[1:]:
        r.text = ""


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("rewrites")
    ap.add_argument("--out", required=True)
    args = ap.parse_args()

    doc = Document(args.docx)
    with open(args.rewrites) as f:
        rewrites = json.load(f)

    # Normalize keys to int paragraph indices.
    rewrites = {int(k): v for k, v in rewrites.items()}

    paragraphs = doc.paragraphs
    applied = 0
    for idx, new_text in rewrites.items():
        if idx < 0 or idx >= len(paragraphs):
            print(f"WARN: id {idx} out of range, skipping", file=sys.stderr)
            continue
        before = paragraphs[idx].text.strip()
        set_paragraph_text(paragraphs[idx], new_text)
        after = paragraphs[idx].text.strip()
        applied += 1
        print(f"  #{idx}\n    - {before[:90]}\n    + {after[:90]}", file=sys.stderr)

    doc.save(args.out)
    print(f"Applied {applied} rewrite(s) -> {args.out}", file=sys.stderr)


if __name__ == "__main__":
    main()
